import os
import logging
from typing import Optional, Tuple
from pathlib import Path

# Библиотеки для работы с файлами
try:
    import pdfplumber  # Более надёжный, чем PyPDF2
except ImportError:
    pdfplumber = None

try:
    from docx import Document
except ImportError:
    Document = None

from config import logger, MAX_TOKENS

# ----------------------------------------------------------------------
# Класс для извлечения текста из файлов
# ----------------------------------------------------------------------
class FileTextExtractor:
    """Извлекает текст из PDF и DOCX файлов."""
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.txt'}
    
    @classmethod
    def extract(cls, file_path: str) -> Tuple[str, str]:
        """
        Извлекает текст из файла.
        
        Args:
            file_path: Путь к файлу
            
        Returns:
            Tuple[text_content, file_type]
            
        Raises:
            ValueError: Если формат файла не поддерживается
            FileNotFoundError: Если файл не найден
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"Файл не найден: {file_path}")
        
        ext = path.suffix.lower()
        
        if ext not in cls.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Неподдерживаемый формат: {ext}. "
                           f"Допустимые: {cls.SUPPORTED_EXTENSIONS}")
        
        if ext == '.pdf':
            return cls._extract_pdf(path), 'pdf'
        elif ext == '.docx':
            return cls._extract_docx(path), 'docx'
        elif ext == '.txt':
            return cls._extract_txt(path), 'txt'
        
        raise ValueError(f"Неизвестный формат: {ext}")
    
    @staticmethod
    def _extract_pdf(path: Path) -> str:
        """Извлекает текст из PDF."""
        if pdfplumber is None:
            raise ImportError("Установите pdfplumber: pip install pdfplumber")
        
        text_parts = []
        
        try:
            with pdfplumber.open(path) as pdf:
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"[Страница {i + 1}]\n{page_text}")
        except Exception as e:
            logger.error(f"Ошибка чтения PDF {path}: {e}")
            raise RuntimeError(f"Не удалось прочитать PDF: {e}")
        
        return '\n\n'.join(text_parts)
    
    @staticmethod
    def _extract_docx(path: Path) -> str:
        """Извлекает текст из DOCX."""
        if Document is None:
            raise ImportError("Установите python-docx: pip install python-docx")
        
        try:
            doc = Document(path)
            text_parts = []
            
            # Извлекаем текст из параграфов
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Дополнительно: текст из таблиц
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
        except Exception as e:
            logger.error(f"Ошибка чтения DOCX {path}: {e}")
            raise RuntimeError(f"Не удалось прочитать DOCX: {e}")
        
        return '\n\n'.join(text_parts)
    
    @staticmethod
    def _extract_txt(path: Path) -> str:
        """Извлекает текст из TXT."""
        try:
            # Пробуем разные кодировки
            encodings = ['utf-8', 'cp1251', 'latin-1']
            
            for enc in encodings:
                try:
                    return path.read_text(encoding=enc)
                except UnicodeDecodeError:
                    continue
            
            raise UnicodeDecodeError("Не удалось определить кодировку файла")
            
        except Exception as e:
            logger.error(f"Ошибка чтения TXT {path}: {e}")
            raise RuntimeError(f"Не удалось прочитать TXT: {e}")


# ----------------------------------------------------------------------
# Интеграция с LLM (на основе вашего существующего кода)
# ----------------------------------------------------------------------
async def process_file_to_slides(
    file_path: str,
    user_prompt: str = "",
    max_slides: int = 6,
    tone: str = "formal",
    image_service: str = "sd",
) -> dict:
    """
    Полный пайплайн: файл → текст → LLM → структура слайдов.
    
    Args:
        file_path: Путь к файлу (PDF/DOCX/TXT)
        user_prompt: Дополнительный промпт от пользователя
        max_slides: Максимальное количество слайдов
        tone: Тон повествования
        image_service: Сервис для генерации изображений (sd/yaArt)
        
    Returns:
        Dict со слайдами и метаданными
    """
    from llm_request import summarize_to_structure  # ваш существующий модуль
    
    # 1. Извлекаем текст
    logger.info(f"Чтение файла: {file_path}")
    text_content, file_type = FileTextExtractor.extract(file_path)
    
    if not text_content.strip():
        raise ValueError("Файл пуст или не содержит извлекаемого текста")
    
    logger.info(f"Извлечено {len(text_content)} символов из {file_type}")
    
    # 2. Отправляем в LLM
    logger.info("Отправка текста в LLM...")
    slides_structure = await summarize_to_structure(
        raw_text=text_content,
        user_prompt=user_prompt,
        max_slides=max_slides,
        tone=tone,
    )
    
    # 3. Формируем ответ
    result = {
        "file_type": file_type,
        "original_length": len(text_content),
        "slides_count": len(slides_structure),
        "slides": slides_structure,
        "image_service": image_service,
    }
    
    logger.info(f"Генерация завершена: {len(slides_structure)} слайдов")
    return result


# ----------------------------------------------------------------------
# FastAPI endpoint (для интеграции с вашим бэкендом)
# ----------------------------------------------------------------------
from fastapi import UploadFile, HTTPException
import tempfile
import shutil

async def handle_file_upload(
    file: UploadFile,
    prompt: str = "",
    slides: int = 5,
    style: str = "corporate",
    tone: str = "formal",
    image_service: str = "sd",
) -> dict:
    """
    Обработчик загрузки файла для FastAPI endpoint.
    
    Используется в вашем API: POST /api/generate
    """
    # Валидация расширения
    allowed_extensions = {'.pdf', '.docx', '.txt'}
    file_ext = Path(file.filename).suffix.lower()
    
    if file_ext not in allowed_extensions:
        raise HTTPException(
            status_code=400,
            detail=f"Неподдерживаемый формат. Допустимые: {allowed_extensions}"
        )
    
    # Временное сохранение файла
    temp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as tmp:
            shutil.copyfileobj(file.file, tmp)
            temp_path = tmp.name
        
        # Обработка
        result = await process_file_to_slides(
            file_path=temp_path,
            user_prompt=prompt,
            max_slides=slides,
            tone=tone,
            image_service=image_service,
        )
        
        return result
        
    except FileNotFoundError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.exception(f"Критическая ошибка обработки файла: {e}")
        raise HTTPException(status_code=500, detail=f"Внутренняя ошибка: {str(e)}")
    finally:
        # Очистка временного файла
        if temp_path and os.path.exists(temp_path):
            os.unlink(temp_path)


# ----------------------------------------------------------------------
# Пример использования (для тестирования)
# ----------------------------------------------------------------------
if __name__ == "__main__":
    import asyncio
    
    async def test():
        # Тест с PDF
        try:
            result = await process_file_to_slides(
                file_path="test_document.pdf",
                user_prompt="Создай презентацию по этому документу",
                max_slides=5,
                tone="formal",
            )
            print(f"✅ Успешно: {result['slides_count']} слайдов")
            for slide in result['slides']:
                print(f"  • {slide['title']}")
        except Exception as e:
            print(f"❌ Ошибка: {e}")
    
    asyncio.run(test())